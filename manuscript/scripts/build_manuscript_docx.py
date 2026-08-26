from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
SOURCE = ROOT / "manuscript" / "Limits_of_Task_Set_Reduction_manuscript.md"
OUTPUT = ROOT / "manuscript" / "Limits_of_Task_Set_Reduction_EMSE_draft_v3.1.docx"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "20262E"
MUTED = "66717C"
TABLE_FILL = "F4F6F9"
GRID = "C8CED5"
GOLD = "9A7418"
CONTENT_DXA = 9360
TABLE_INDENT_DXA = 120


def set_font(run, name="Calibri", size=None, color=None, bold=None, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in("w:tcMar")
    if tcMar is None:
        tcMar = OxmlElement("w:tcMar")
        tcPr.append(tcMar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tcMar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tcMar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    trPr = row._tr.get_or_add_trPr()
    tblHeader = OxmlElement("w:tblHeader")
    tblHeader.set(qn("w:val"), "true")
    trPr.append(tblHeader)


def set_table_borders(table):
    tblPr = table._tbl.tblPr
    borders = tblPr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tblPr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), "4")
        node.set(qn("w:color"), GRID)


def set_cell_fill(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tcPr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_table_geometry(table, widths):
    assert sum(widths) == CONTENT_DXA
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tblPr = table._tbl.tblPr
    layout = tblPr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tblPr.append(layout)
    layout.set(qn("w:type"), "fixed")
    tblW = tblPr.find(qn("w:tblW"))
    tblW.set(qn("w:w"), str(CONTENT_DXA))
    tblW.set(qn("w:type"), "dxa")
    tblInd = tblPr.find(qn("w:tblInd"))
    if tblInd is None:
        tblInd = OxmlElement("w:tblInd")
        tblPr.append(tblInd)
    tblInd.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tblInd.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tcPr = cell._tc.get_or_add_tcPr()
            tcW = tcPr.find(qn("w:tcW"))
            if tcW is None:
                tcW = OxmlElement("w:tcW")
                tcPr.append(tcW)
            tcW.set(qn("w:w"), str(widths[idx]))
            tcW.set(qn("w:type"), "dxa")
            cell.width = Inches(widths[idx] / 1440)
            set_cell_margins(cell)


def column_widths(col_count):
    patterns = {
        4: [3000, 2120, 2120, 2120],
        5: [1850, 1200, 2300, 2300, 1710],
        7: [1750, 900, 900, 1000, 1760, 1550, 1500],
    }
    if col_count in patterns:
        return patterns[col_count]
    base = CONTENT_DXA // col_count
    widths = [base] * col_count
    widths[-1] += CONTENT_DXA - sum(widths)
    return widths


def add_page_field(paragraph):
    run = paragraph.add_run()
    set_font(run, size=8.5, color=MUTED)

    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    run._r.append(begin)

    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    run._r.append(instruction)

    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    run._r.append(separate)

    placeholder = OxmlElement("w:t")
    placeholder.text = "1"
    run._r.append(placeholder)

    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(end)


def configure_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    pf = normal.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.space_before = Pt(0)
    pf.space_after = Pt(8)
    pf.line_spacing = 1.333

    for style_name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ):
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.keep_together = True

    caption = doc.styles["Caption"]
    caption.font.name = "Calibri"
    caption._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    caption._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    caption.font.size = Pt(9.5)
    caption.font.italic = False
    caption.font.color.rgb = RGBColor.from_string(INK)
    caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    caption.paragraph_format.space_before = Pt(4)
    caption.paragraph_format.space_after = Pt(4)
    caption.paragraph_format.keep_with_next = True

    if "References" not in [s.name for s in doc.styles]:
        ref = doc.styles.add_style("References", WD_STYLE_TYPE.PARAGRAPH)
    else:
        ref = doc.styles["References"]
    ref.font.name = "Calibri"
    ref._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    ref._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    ref.font.size = Pt(9.5)
    ref.font.color.rgb = RGBColor.from_string(INK)
    ref.paragraph_format.left_indent = Inches(0.25)
    ref.paragraph_format.first_line_indent = Inches(-0.25)
    ref.paragraph_format.space_after = Pt(5)
    ref.paragraph_format.line_spacing = 1.0

    for style_name in ("List Bullet", "List Number"):
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.194)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.208


def configure_sections(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = True

    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("Limits of Task-Set Reduction in SWE-bench Verified")
    set_font(r, size=8.5, color=MUTED, bold=True)
    r = p.add_run("  |  Manuscript")
    set_font(r, size=8.5, color=MUTED)

    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_before = Pt(0)
    r = p.add_run("Page ")
    set_font(r, size=8.5, color=MUTED)
    add_page_field(p)

    first_footer = section.first_page_footer
    p = first_footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Manuscript for internal review")
    set_font(r, size=8.5, color=MUTED)


def add_inline(paragraph, text):
    token_re = re.compile(r"(τ_b|https?://\S+|`[^`]+`|\*\*[^*]+\*\*|\*[^*]+\*)")
    pos = 0
    for match in token_re.finditer(text):
        if match.start() > pos:
            paragraph.add_run(text[pos:match.start()])
        token = match.group(0)
        if token == "τ_b":
            paragraph.add_run("τ")
            run = paragraph.add_run("b")
            run.font.subscript = True
        elif token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
        elif token.startswith("*"):
            run = paragraph.add_run(token[1:-1])
            run.italic = True
        elif token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            set_font(run, name="Consolas", size=9.5, color=DARK_BLUE)
        else:
            run = paragraph.add_run(token)
            run.font.color.rgb = RGBColor.from_string(BLUE)
            run.font.underline = True
        pos = match.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


def add_cover(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(68)
    p.add_run("")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(18)
    r = p.add_run("RESEARCH ARTICLE")
    set_font(r, size=10.5, color=GOLD, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(12)
    r = p.add_run("Limits of Task-Set Reduction in SWE-bench Verified")
    set_font(r, size=24, color=INK, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(34)
    r = p.add_run("A Temporal Study of Leaderboard Ranking Reliability")
    set_font(r, size=15, color=DARK_BLUE)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run("Anonymous manuscript for review")
    set_font(r, size=11, color=INK, bold=True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(34)
    r = p.add_run("Target journal: Empirical Software Engineering")
    set_font(r, size=10.5, color=MUTED, italic=True)

    doc.add_page_break()


def is_table_line(line):
    return line.startswith("|") and line.endswith("|")


def parse_table(lines, start):
    block = []
    i = start
    while i < len(lines) and is_table_line(lines[i].strip()):
        block.append(lines[i].strip())
        i += 1
    rows = []
    for line in block:
        cells = [c.strip() for c in line.strip("|").split("|")]
        if all(re.fullmatch(r":?-{3,}:?", c) for c in cells):
            continue
        rows.append(cells)
    return rows, i


def add_table(doc, rows):
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    set_table_geometry(table, column_widths(len(rows[0])))
    set_table_borders(table)
    set_repeat_table_header(table.rows[0])
    for i, data_row in enumerate(rows):
        for j, value in enumerate(data_row):
            cell = table.cell(i, j)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if i == 0:
                set_cell_fill(cell, TABLE_FILL)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if j == 0 else WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.0
            p.clear()
            add_inline(p, value)
            for run in p.runs:
                set_font(run, size=8.5, color=INK, bold=(i == 0))
    after = doc.add_paragraph()
    after.paragraph_format.space_before = Pt(0)
    after.paragraph_format.space_after = Pt(3)
    return table


def add_image(doc, rel_path, alt_text):
    path = (SOURCE.parent / rel_path).resolve()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run()
    shape = run.add_picture(str(path), width=Inches(6.35))
    doc_pr = shape._inline.docPr
    doc_pr.set("descr", alt_text)
    return p


def collect_paragraph(lines, start):
    parts = []
    i = start
    while i < len(lines):
        line = lines[i].rstrip()
        stripped = line.strip()
        if not stripped:
            break
        if stripped.startswith("#") or is_table_line(stripped) or stripped.startswith("!["):
            break
        if re.match(r"^[-*] ", stripped) or re.match(r"^\d+\. ", stripped):
            break
        parts.append(stripped)
        i += 1
    return " ".join(parts), i


def collect_list_item(lines, start, ordered):
    pattern = r"^\d+\.\s+" if ordered else r"^[-*]\s+"
    first = re.sub(pattern, "", lines[start].strip())
    parts = [first]
    i = start + 1
    while i < len(lines):
        raw = lines[i]
        stripped = raw.strip()
        if not stripped:
            break
        if re.match(r"^[-*] ", stripped) or re.match(r"^\d+\. ", stripped):
            break
        if stripped.startswith("#") or is_table_line(stripped) or stripped.startswith("!["):
            break
        parts.append(stripped)
        i += 1
    return " ".join(parts), i


def build():
    doc = Document()
    configure_styles(doc)
    configure_sections(doc)
    add_cover(doc)

    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    i = 5  # skip source title and draft metadata; cover already contains them
    in_references = False
    while i < len(lines):
        stripped = lines[i].strip()
        if not stripped:
            i += 1
            continue
        if stripped.startswith("<!--") and stripped.endswith("-->"):
            i += 1
            continue
        if stripped.startswith("## "):
            heading = stripped[3:]
            if heading == "References":
                in_references = True
            p = doc.add_paragraph(style="Heading 1")
            add_inline(p, heading)
            i += 1
            continue
        if stripped.startswith("### "):
            p = doc.add_paragraph(style="Heading 2")
            add_inline(p, stripped[4:])
            i += 1
            continue
        if stripped.startswith("#### "):
            p = doc.add_paragraph(style="Heading 3")
            add_inline(p, stripped[5:])
            i += 1
            continue
        if is_table_line(stripped):
            rows, i = parse_table(lines, i)
            add_table(doc, rows)
            continue
        image_match = re.match(r"!\[(.+?)\]\((.+?)\)", stripped)
        if image_match:
            add_image(doc, image_match.group(2), image_match.group(1))
            i += 1
            continue
        if re.match(r"^\d+\. ", stripped) and not in_references:
            number = re.match(r"^(\d+)\. ", stripped).group(1)
            text, i = collect_list_item(lines, i, True)
            p = doc.add_paragraph(style="Normal")
            p.paragraph_format.left_indent = Inches(0.28)
            p.paragraph_format.first_line_indent = Inches(-0.20)
            p.paragraph_format.space_after = Pt(5)
            p.add_run(f"{number}. ")
            add_inline(p, text)
            continue
        if re.match(r"^[-*] ", stripped) and not in_references:
            text, i = collect_list_item(lines, i, False)
            p = doc.add_paragraph(style="List Bullet")
            add_inline(p, text)
            continue

        text, i = collect_paragraph(lines, i)
        style = "References" if in_references else "Normal"
        if text.startswith("**Table ") or text.startswith("**Fig. "):
            style = "Caption"
        p = doc.add_paragraph(style=style)
        # Keep the two longer decision tables intact. Without a page break,
        # LibreOffice places their captions at the foot of a page and splits
        # the body across pages, which makes the decision rows hard to audit.
        if (
            text.startswith("**Table 4.")
            or text.startswith("**Table 8.")
        ):
            p.paragraph_format.page_break_before = True
        if text.startswith("**Keywords:"):
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        if text.startswith("The formal GitHub Actions execution"):
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        add_inline(p, text)

    props = doc.core_properties
    props.title = "Limits of Task-Set Reduction in SWE-bench Verified: A Temporal Study of Leaderboard Ranking Reliability"
    props.subject = "Temporal ranking reliability of reduced coding-agent benchmark task sets"
    props.author = "Anonymous"
    # Non-breaking spaces keep each multiword keyword intact when LibreOffice
    # exports the DOCX keyword list to PDF metadata.
    props.keywords = "coding\u00a0agents; benchmark\u00a0reduction; leaderboard\u00a0reliability; SWE-bench\u00a0Verified; temporal\u00a0validation; ranking\u00a0uncertainty"
    props.comments = "Temporal ranking-reliability study"
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
